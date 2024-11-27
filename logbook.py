import os
from datetime import datetime
from docx import Document

# Bestandspad voor het logboek (Word-bestand)
logbestand = "logboek.docx"

# Functie om een nieuw logbericht toe te voegen aan het Word-bestand
def log_update(bericht):
    # Krijg de huidige datum en tijd
    huidige_tijd = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Maak het document of open het bestaande document
    if os.path.exists(logbestand):
        document = Document(logbestand)
    else:
        document = Document()
        document.add_heading('Logboek', 0)
    
    # Voeg het nieuwe logbericht toe
    document.add_paragraph(f"{huidige_tijd} - {bericht}")
    
    # Sla het document op
    document.save(logbestand)
    
    print("Logbericht toegevoegd!")

# Functie om het logboek in te zien
def toon_logboek():
    if os.path.exists(logbestand):
        document = Document(logbestand)
        for para in document.paragraphs:
            print(para.text)
    else:
        print("Logboek bestaat nog niet.")

# Hoofdmenu van de app
def main():
    while True:
        print("\nLogboek App")
        print("1. Logbericht toevoegen")
        print("2. Logboek tonen")
        print("3. Afsluiten")
        
        keuze = input("Maak een keuze (1/2/3): ")
        
        if keuze == '1':
            bericht = input("Voer het logbericht in: ")
            log_update(bericht)
        elif keuze == '2':
            toon_logboek()
        elif keuze == '3':
            print("App wordt afgesloten.")
            break
        else:
            print("Ongeldige keuze, probeer opnieuw.")

# Start de applicatie
if __name__ == "__main__":
    main()
